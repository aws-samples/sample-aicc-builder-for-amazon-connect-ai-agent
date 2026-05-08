"""
AICC Builder Agent — ECS Fargate Entrypoint

FastAPI application replacing BedrockAgentCoreApp for ECS Fargate deployment.
Provides the same WebSocket + HTTP interface as the AgentCore agent.py.

Endpoints (port 8080):
- GET  /ping         — ALB health check
- POST /invocations  — HTTP invocations (sync)
- WS   /ws           — WebSocket bidirectional streaming

Key differences from AgentCore mode:
- No BedrockAgentCoreApp / RequestContext — session_id via query param
- No AgentCore Memory — conversation history persisted to S3 Files NFS
- No Redis — 3-tier session storage (memory → NFS → DynamoDB metadata)
- Cognito JWT validation for WebSocket auth
- SIGTERM handler for graceful shutdown with S3 Files flush
- Custom CloudWatch metric publishing (ActiveWebSocketConnections)
"""

import os
import re
import sys
import json
import signal
import asyncio
import logging
import time
import traceback
from typing import Optional, Dict, Any
from datetime import datetime

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, Query, HTTPException, Depends, Header
from fastapi.responses import JSONResponse
from pydantic import BaseModel

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))

from strands import Agent
from strands.models import BedrockModel
from botocore.config import Config as BotocoreConfig

# Import shared modules from src/
from prompts.system_prompt import SYSTEM_PROMPT, get_document_analysis_prompt, get_phase_system_prompt
from context import format_history_for_injection
from context.modification_tracking import (
    format_modification_state_block,
    record_modification_request,
    suggest_placeholder,
)

# Import utility tools
from tools import (
    introspect_database,
    save_operation_spec,
    get_operation_spec,
    list_operations,
    get_all_operation_ids,
    get_all_tool_ids,
    update_operation_spec,
    format_operation_summary,
    save_session_flow_config,
    get_session_flow_config_tool,
    save_infrastructure_spec,
    get_infrastructure_spec_tool,
    infer_missing_tools,
    stream_fallback_asset,
    merge_infrastructure_fragments,
    merge_openapi_fragments,
    asset_lookup,
    validate_parameter_consistency,
    read_workspace_file,
    write_workspace_file,
    append_workspace_file,
    list_workspace_dir,
    patch_workspace_file,
    find_workspace_files,
    grep_workspace,
)
from tools.project_workspace import (
    save_requirement_document,
    load_requirement_document,
)
from tools.interview_completion import complete_interview, check_interview_handoff
from tools.streaming_callback import set_session_id as set_streaming_session_id, set_message_index

# Import Sub-Agent tools
from agents import (
    research_agent,
    faq_generator_agent,
    lambda_generator_agent,
    openapi_generator_agent,
    prompt_generator_agent,
    contact_flow_generator_agent,
    infrastructure_generator_agent,
    reviewer_agent,
)

# Import Sub-Agent callback handler setters
from agents.research_agent.agent import set_callback_handler as set_research_callback, clear_session as clear_research_session
from agents.faq_generator.agent import set_callback_handler as set_faq_callback, clear_session as clear_faq_session
from agents.lambda_generator.agent import set_callback_handler as set_lambda_callback
from agents.openapi_generator.agent import set_callback_handler as set_openapi_callback
from agents.prompt_generator.agent import set_callback_handler as set_prompt_callback
from agents.contact_flow_generator.agent import set_callback_handler as set_contact_flow_callback
from agents.infrastructure_generator.agent import set_callback_handler as set_infrastructure_callback
from agents.reviewer_agent.agent import set_callback_handler as set_reviewer_callback

# Prompt hot-reload
try:
    from prompts.prompt_loader import get_system_prompt
except ImportError:
    get_system_prompt = None

# S3 Files context store
from context.s3files_store import S3FilesContextStore
from context.message_log import get_message_log
from context.generation_progress import update_from_new_messages as _update_generation_progress
from context.generation_progress import read_progress as _read_generation_progress
from context.generation_progress import record_tool_completion as _record_tool_completion
from context.generation_progress import detect_phase as _detect_phase
from context.generation_progress import read_phase as _read_phase
from context.generation_progress import update_phase as _update_phase
from context.generation_progress import get_frontend_progress_state as _get_frontend_progress

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger("aicc-ecs")

# ========================================
# Configuration
# ========================================
S3FILES_MOUNT = os.environ.get("S3FILES_MOUNT_PATH", "/mnt/s3")
AWS_REGION = os.environ.get("AWS_REGION", "ap-northeast-2")
ASSETS_BUCKET_NAME = os.environ.get("ASSETS_BUCKET_NAME", "")

# ========================================
# Auto Progress Update Mapping (same as AgentCore)
# ========================================
SUBAGENT_TO_PROGRESS_ID = {
    "infrastructure_generator_agent": "cdk",
    "merge_infrastructure_fragments": "cdk",
    "merge_openapi_fragments": "openapi",
    "lambda_generator_agent": "lambda",
    "openapi_generator_agent": "openapi",
    "prompt_generator_agent": "prompt",
    "contact_flow_generator_agent": "contact_flow",
    "faq_generator_agent": "knowledge_base",
    "reviewer_agent": "review",
}

# ========================================
# Phase-Specific Tool Lists
# ========================================
INTERVIEW_TOOLS = [
    introspect_database,
    save_operation_spec,
    get_operation_spec,
    list_operations,
    update_operation_spec,
    format_operation_summary,
    save_session_flow_config,
    get_session_flow_config_tool,
    save_infrastructure_spec,
    get_infrastructure_spec_tool,
    infer_missing_tools,
    save_requirement_document,
    load_requirement_document,
    # NFS workspace file tools
    read_workspace_file,
    write_workspace_file,
    append_workspace_file,
    list_workspace_dir,
    patch_workspace_file,
    find_workspace_files,
    grep_workspace,
    # Sub-agents available during interview
    research_agent,
    # Interview completion signal
    complete_interview,
]

GENERATION_TOOLS = [
    get_operation_spec,
    list_operations,
    get_all_operation_ids,
    get_all_tool_ids,
    get_session_flow_config_tool,
    get_infrastructure_spec_tool,
    load_requirement_document,
    # NFS workspace file tools
    read_workspace_file,
    write_workspace_file,
    append_workspace_file,
    list_workspace_dir,
    patch_workspace_file,
    find_workspace_files,
    grep_workspace,
    # Generation utilities
    stream_fallback_asset,
    merge_infrastructure_fragments,
    merge_openapi_fragments,
    asset_lookup,
    validate_parameter_consistency,
    # Sub-agents for generation
    lambda_generator_agent,
    openapi_generator_agent,
    prompt_generator_agent,
    contact_flow_generator_agent,
    infrastructure_generator_agent,
    reviewer_agent,
    faq_generator_agent,
    research_agent,
]


def get_tools_for_phase(phase: str) -> list:
    """Return the appropriate tool list for the given phase."""
    if phase == "interview":
        return INTERVIEW_TOOLS
    return GENERATION_TOOLS


# ========================================
# FastAPI Application
# ========================================
app = FastAPI(title="AICC Builder ECS", version="2.0.0")

# ========================================
# Session & Connection Tracking
# ========================================
session_store: Dict[str, Dict[str, Any]] = {}
_active_ws_connections: set = set()
_context_store = S3FilesContextStore(S3FILES_MOUNT)

# ========================================
# Background Task Registry
# ========================================
# Tracks running agent tasks per session so they survive WebSocket disconnects.
# Key: session_id, Value: {"task": asyncio.Task, "ws_ref": WebSocket|None, "started_at": float}
_background_tasks: Dict[str, Dict[str, Any]] = {}
_background_tasks_lock: Optional[asyncio.Lock] = None  # created in startup()

# CloudWatch metric publishing
_cw_client = None

def _get_cw_client():
    global _cw_client
    if _cw_client is None:
        import boto3
        _cw_client = boto3.client("cloudwatch", region_name=AWS_REGION)
    return _cw_client

async def publish_ws_metric():
    """Publish ActiveWebSocketConnections metric for auto-scaling."""
    try:
        cw = _get_cw_client()
        cw.put_metric_data(
            Namespace="AiccBuilder/ECS",
            MetricData=[{
                "MetricName": "ActiveWebSocketConnections",
                "Value": len(_active_ws_connections),
                "Unit": "Count",
            }],
        )
    except Exception as e:
        logger.warning(f"CloudWatch metric publish failed: {e}")

# Background metric publisher
_metric_task: Optional[asyncio.Task] = None

async def _metric_publisher():
    while True:
        await publish_ws_metric()
        await asyncio.sleep(60)

@app.on_event("startup")
async def startup():
    global _metric_task, _background_tasks_lock
    _background_tasks_lock = asyncio.Lock()
    _metric_task = asyncio.create_task(_metric_publisher())
    logger.info(f"AICC Builder ECS started. S3FILES_MOUNT={S3FILES_MOUNT}, REGION={AWS_REGION}")

@app.on_event("shutdown")
async def shutdown():
    if _metric_task:
        _metric_task.cancel()
    await _graceful_shutdown()

# ========================================
# SIGTERM Graceful Shutdown
# ========================================
async def _graceful_shutdown():
    """Flush active sessions to S3 Files and close WebSocket connections."""
    logger.info(f"Graceful shutdown: flushing {len(session_store)} sessions, closing {len(_active_ws_connections)} WS connections")

    for session_id, session in session_store.items():
        try:
            # Save conversation history to NFS
            history = session.get("conversation_history", [])
            if history:
                _context_store.save_conversation_history(session_id, history)
            # Save session context
            ctx = _context_store.get_session(session_id)
            if ctx:
                _context_store.save_session(ctx)
            logger.info(f"Flushed session {session_id}")
        except Exception as e:
            logger.error(f"Failed to flush session {session_id}: {e}")

def _sigterm_handler(signum, frame):
    logger.info("SIGTERM received, initiating graceful shutdown")
    loop = asyncio.get_event_loop()
    loop.create_task(_graceful_shutdown())

signal.signal(signal.SIGTERM, _sigterm_handler)

# ========================================
# Cognito JWT Validation
# ========================================
_jwks_cache = None

async def validate_cognito_token(token: str) -> Optional[Dict]:
    """Validate Cognito JWT token. Returns claims dict or None."""
    if not token:
        return None

    user_pool_id = os.environ.get("USER_POOL_ID", "")
    if not user_pool_id:
        # If no user pool configured, skip validation (local dev)
        logger.warning("USER_POOL_ID not set, skipping JWT validation")
        return {"sub": "local-dev"}

    try:
        from jose import jwt, JWTError
        import urllib.request

        global _jwks_cache
        if _jwks_cache is None:
            jwks_url = f"https://cognito-idp.{AWS_REGION}.amazonaws.com/{user_pool_id}/.well-known/jwks.json"
            with urllib.request.urlopen(jwks_url, timeout=5) as resp:
                _jwks_cache = json.loads(resp.read())

        # Decode without verification first to get kid
        unverified = jwt.get_unverified_header(token)
        kid = unverified.get("kid")

        # Find matching key
        key = None
        for k in _jwks_cache.get("keys", []):
            if k.get("kid") == kid:
                key = k
                break

        if not key:
            logger.warning("JWT kid not found in JWKS")
            return None

        claims = jwt.decode(
            token,
            key,
            algorithms=["RS256"],
            audience=os.environ.get("USER_POOL_CLIENT_ID", ""),
            issuer=f"https://cognito-idp.{AWS_REGION}.amazonaws.com/{user_pool_id}",
        )
        return claims
    except Exception as e:
        logger.warning(f"JWT validation failed: {e}")
        return None

# ========================================
# Model Configuration (same as AgentCore)
# ========================================
def get_model_config():
    model_id = os.environ.get(
        "BEDROCK_MODEL_ID",
        "global.anthropic.claude-opus-4-6-v1"
    )
    return SafeBedrockModel(
        model_id=model_id,
        region_name=AWS_REGION,
        boto_client_config=BotocoreConfig(
            read_timeout=300,
            retries={"max_attempts": 3, "mode": "adaptive"},
        )
    )


class SafeBedrockModel(BedrockModel):
    """Layer 4 defense: sanitize messages right before Bedrock API call.

    Strands SDK manages agent.messages during active conversation and may
    inject toolResult blocks that mismatch with toolUse blocks.  This wrapper
    intercepts _format_request (the INTERNAL method called by _stream) to fix
    pairing issues just before the wire call.

    NOTE: The SDK calls self._format_request() (underscore prefix), NOT the
    public format_request().  We must override _format_request to ensure our
    sanitization actually runs.
    """

    def _format_request(self, messages, tool_specs=None, system_prompt_content=None, tool_choice=None):
        """Override internal _format_request to sanitize messages before Bedrock API call."""
        messages = self._fix_messages_for_bedrock(messages)
        return super()._format_request(messages, tool_specs=tool_specs, system_prompt_content=system_prompt_content, tool_choice=tool_choice)

    @staticmethod
    def _fix_messages_for_bedrock(messages: list) -> list:
        """Ensure messages satisfy Bedrock ConverseStream constraints:
        1. First message must be role=user
        2. Strict user/assistant alternation
        3. Every toolResult has matching toolUse in preceding assistant
        4. toolResult count <= toolUse count

        This runs on EVERY Bedrock API call during agent execution, not just
        on the initial history load.  Critical for catching issues that arise
        mid-conversation as the SDK mutates agent.messages.
        """
        if not messages:
            return messages

        # 0. Filter out system-role messages (Bedrock only accepts user/assistant)
        fixed = [m for m in messages if m.get("role") in ("user", "assistant")]

        # 1. Drop leading non-user messages
        while fixed and fixed[0].get("role") != "user":
            logger.warning("[SafeBedrockModel] Dropping leading non-user message")
            fixed.pop(0)

        if not fixed:
            return fixed

        # 2. Enforce role alternation — merge consecutive same-role messages
        merged = [fixed[0]]
        for msg in fixed[1:]:
            if msg.get("role") == merged[-1].get("role"):
                # Merge content blocks
                prev_content = merged[-1].get("content", [])
                curr_content = msg.get("content", [])
                if isinstance(prev_content, list) and isinstance(curr_content, list):
                    merged[-1] = {"role": msg["role"], "content": prev_content + curr_content}
                # else: skip malformed
            else:
                merged.append(msg)
        fixed = merged

        # 3. Fix toolResult/toolUse pairing
        for i in range(len(fixed)):
            msg = fixed[i]
            content = msg.get("content")
            if not isinstance(content, list) or msg.get("role") != "user":
                continue

            tool_result_ids = {
                b["toolResult"]["toolUseId"]
                for b in content
                if isinstance(b, dict) and "toolResult" in b and b["toolResult"].get("toolUseId")
            }
            if not tool_result_ids:
                continue

            # Collect toolUse IDs from preceding assistant
            preceding_tool_use_ids = set()
            if i > 0 and fixed[i - 1].get("role") == "assistant":
                prev_content = fixed[i - 1].get("content", [])
                if isinstance(prev_content, list):
                    preceding_tool_use_ids = {
                        b["toolUse"]["toolUseId"]
                        for b in prev_content
                        if isinstance(b, dict) and "toolUse" in b and b["toolUse"].get("toolUseId")
                    }

            excess = tool_result_ids - preceding_tool_use_ids
            if excess:
                logger.warning(f"[SafeBedrockModel] Removing {len(excess)} excess toolResult blocks at msg {i}")
                cleaned = [
                    b for b in content
                    if not (isinstance(b, dict) and "toolResult" in b and b["toolResult"].get("toolUseId") in excess)
                ]
                if cleaned:
                    fixed[i] = {"role": "user", "content": cleaned}
                else:
                    fixed[i] = {"role": "user", "content": [{"text": "(tool results removed)"}]}

        # 4. Remove trailing assistant toolUse without following toolResult
        if fixed and fixed[-1].get("role") == "assistant":
            last_content = fixed[-1].get("content", [])
            if isinstance(last_content, list):
                has_tool_use = any(isinstance(b, dict) and "toolUse" in b for b in last_content)
                if has_tool_use:
                    cleaned = [b for b in last_content if not (isinstance(b, dict) and "toolUse" in b)]
                    if cleaned:
                        fixed[-1] = {"role": "assistant", "content": cleaned}
                    else:
                        fixed.pop()
                        logger.warning("[SafeBedrockModel] Removed trailing assistant with only toolUse blocks")

        return fixed

# ========================================
# Session Management
# ========================================
def get_or_create_session(session_id: str) -> Dict[str, Any]:
    """Get or create session, restoring from NFS if available."""
    if session_id in session_store:
        session_store[session_id]["last_active"] = datetime.utcnow().isoformat()
        return session_store[session_id]

    model = get_model_config()

    # Use hot-reloaded prompt if available
    system_prompt = SYSTEM_PROMPT
    if get_system_prompt is not None:
        try:
            reloaded = get_system_prompt()
            if reloaded:
                system_prompt = reloaded
        except Exception:
            pass

    # Determine phase to select appropriate tools
    initial_phase = _detect_phase(session_id)
    tools = get_tools_for_phase(initial_phase)

    agent = Agent(
        model=model,
        system_prompt=system_prompt,
        tools=tools,
    )

    # Restore conversation history from NFS
    conversation_history = []
    try:
        restored = _context_store.load_conversation_history(session_id)
        if restored:
            conversation_history = _validate_tool_pairs(restored)
            logger.info(f"Restored {len(conversation_history)} messages for session {session_id}")
    except Exception as e:
        logger.warning(f"History restore failed for {session_id}: {e}")

    session_store[session_id] = {
        "agent": agent,
        "model": model,
        "tools": tools,
        "conversation_history": conversation_history,
        "session_data": {},
        "document_mode": False,
        "uploaded_document": None,
        "created_at": datetime.utcnow().isoformat(),
        "last_active": datetime.utcnow().isoformat(),
    }

    logger.info(f"Session created: {session_id}")
    return session_store[session_id]

def summarize_response_for_history(response_text: str, max_length: int = 2000) -> str:
    """Summarize response for history storage."""
    if len(response_text) <= max_length:
        return response_text
    return response_text[:max_length] + f"\n... [truncated from {len(response_text)} chars]"


# ========================================
# Conversation History — Full Strands Format
# ========================================
MAX_HISTORY_MESSAGES = 60  # 30 turns (user + assistant)
MAX_TOOL_RESULT_LENGTH = 1500  # Truncate large tool results
MAX_TOOL_INPUT_LENGTH = 1000   # Truncate large tool inputs


def _extract_new_messages(all_messages: list, prev_count: int) -> list:
    """Extract new messages added by streaming_agent beyond what we loaded from history.

    Strands Agent stores messages as:
      [{"role": "user"/"assistant", "content": [ContentBlock, ...]}]

    ContentBlock can be: {"text": "..."}, {"toolUse": {...}}, {"toolResult": {...}},
    {"image": {...}}, {"reasoningContent": {...}}, etc.

    We keep toolUse/toolResult blocks but truncate large payloads to manage storage size.
    We strip binary content (images, documents) to avoid bloating history.
    """
    new_messages = all_messages[prev_count:]
    sanitized = []
    for msg in new_messages:
        role = msg.get("role")
        content = msg.get("content")
        if not role or not isinstance(content, list):
            continue
        clean_blocks = []
        for block in content:
            if not isinstance(block, dict):
                continue
            if "text" in block:
                # Keep text blocks, truncate if very long
                text = block["text"]
                if len(text) > 3000:
                    text = text[:3000] + f"\n... [truncated from {len(block['text'])} chars]"
                clean_blocks.append({"text": text})
            elif "toolUse" in block:
                # Keep toolUse but truncate large inputs
                tu = block["toolUse"]
                clean_tu = {
                    "name": tu.get("name", ""),
                    "toolUseId": tu.get("toolUseId", ""),
                    "input": _truncate_tool_payload(tu.get("input"), MAX_TOOL_INPUT_LENGTH),
                }
                clean_blocks.append({"toolUse": clean_tu})
            elif "toolResult" in block:
                # Keep toolResult but truncate large results
                tr = block["toolResult"]
                clean_content = []
                for rc in (tr.get("content") or []):
                    if isinstance(rc, dict):
                        if "text" in rc:
                            text = rc["text"]
                            if len(text) > MAX_TOOL_RESULT_LENGTH:
                                text = text[:MAX_TOOL_RESULT_LENGTH] + "... [truncated]"
                            clean_content.append({"text": text})
                        elif "json" in rc:
                            json_str = json.dumps(rc["json"], ensure_ascii=False, default=str)
                            if len(json_str) > MAX_TOOL_RESULT_LENGTH:
                                json_str = json_str[:MAX_TOOL_RESULT_LENGTH] + "... [truncated]"
                            clean_content.append({"text": json_str})
                        # Skip image/document blocks
                clean_tr = {
                    "toolUseId": tr.get("toolUseId", ""),
                    "status": tr.get("status", "success"),
                    "content": clean_content,
                }
                clean_blocks.append({"toolResult": clean_tr})
            # Skip image, document, video, reasoningContent — not needed for memory
        if clean_blocks:
            sanitized.append({"role": role, "content": clean_blocks})

    # Validate: ensure no trailing assistant message with toolUse that lacks
    # a following user message with matching toolResult (partial stream crash).
    if sanitized:
        last = sanitized[-1]
        if last.get("role") == "assistant":
            last_content = last.get("content", [])
            has_tool_use = any(
                isinstance(b, dict) and "toolUse" in b for b in last_content
            )
            if has_tool_use:
                # Trailing assistant with toolUse but no following toolResult.
                # Remove toolUse blocks to prevent mismatch on next turn.
                cleaned = [b for b in last_content if not (isinstance(b, dict) and "toolUse" in b)]
                if cleaned:
                    sanitized[-1] = {"role": "assistant", "content": cleaned}
                else:
                    sanitized.pop()  # remove entirely if only toolUse blocks
                logger.warning("[history] Stripped trailing toolUse from partial stream")

    return sanitized


def _truncate_tool_payload(payload, max_length: int):
    """Truncate tool input/output if its JSON representation is too large."""
    if payload is None:
        return payload
    try:
        s = json.dumps(payload, ensure_ascii=False, default=str)
        if len(s) <= max_length:
            return payload
        # Return a truncated text representation instead
        return {"_truncated": s[:max_length] + "... [truncated]"}
    except Exception:
        return str(payload)[:max_length]


def _find_safe_cut_index(history: list, desired_start: int) -> int:
    """Find the earliest safe index to start slicing history without
    orphaning toolResult blocks.

    Bedrock Converse API requires every toolResult in a user message to
    have a matching toolUse in the immediately preceding assistant message.
    A naive slice can cut between an assistant (toolUse) and user (toolResult),
    leaving orphaned toolResults.

    Strategy:
    1. Forward scan from desired_start, skipping messages with toolResult.
    2. Ensure cut point starts with a user-role message (role alternation).
    3. Skip toolResult-only user messages (they need a preceding assistant).
    """
    if desired_start <= 0:
        return 0
    if desired_start >= len(history):
        return len(history)

    idx = desired_start

    # Forward scan: skip past any toolResult-bearing messages
    while idx < len(history):
        msg = history[idx]
        content = msg.get("content")
        if not isinstance(content, list):
            break  # safe — not a tool message

        has_tool_result = any(
            isinstance(b, dict) and "toolResult" in b for b in content
        )
        if not has_tool_result:
            break  # safe — no orphaned toolResults

        idx += 1

    # Ensure cut point starts with a user-role message
    while idx < len(history):
        msg = history[idx]
        role = msg.get("role")

        if role == "user":
            # Also skip user messages that contain ONLY toolResult blocks
            # (they need a preceding assistant with matching toolUse)
            content = msg.get("content", [])
            if isinstance(content, list):
                non_tool_result = [
                    b for b in content
                    if isinstance(b, dict) and "toolResult" not in b
                ]
                if non_tool_result:
                    break  # Good: user message with real content
                else:
                    idx += 1  # Skip toolResult-only user message
                    continue
            break
        else:
            # Skip assistant messages at cut boundary
            idx += 1

    return min(idx, len(history))


def _validate_tool_pairs(history: list) -> list:
    """Validate and repair toolUse/toolResult pairing in conversation history.

    Multi-pass approach: builds a new list incrementally so that preceding-message
    lookups always reference already-cleaned data.  Repeats until stable (max 10
    passes) to handle cascading orphans where removing one block invalidates another.
    """
    if not history:
        return history

    MAX_PASSES = 10
    current = [m for m in history if m is not None]

    for pass_num in range(MAX_PASSES):
        changed = False
        new_repaired = []

        for i, msg in enumerate(current):
            if msg is None:
                continue
            content = msg.get("content")
            role = msg.get("role")

            if not isinstance(content, list) or not role:
                new_repaired.append(msg)
                continue

            # --- User messages: validate toolResult against preceding assistant ---
            if role == "user":
                tool_result_ids = {
                    b["toolResult"]["toolUseId"]
                    for b in content
                    if isinstance(b, dict) and "toolResult" in b and b["toolResult"].get("toolUseId")
                }
                if tool_result_ids:
                    # Look at the last message in new_repaired (already cleaned)
                    preceding_tool_use_ids = set()
                    if new_repaired and new_repaired[-1].get("role") == "assistant":
                        prev_content = new_repaired[-1].get("content", [])
                        if isinstance(prev_content, list):
                            preceding_tool_use_ids = {
                                b["toolUse"]["toolUseId"]
                                for b in prev_content
                                if isinstance(b, dict) and "toolUse" in b and b["toolUse"].get("toolUseId")
                            }

                    orphaned = tool_result_ids - preceding_tool_use_ids
                    if orphaned:
                        logger.warning(
                            f"[history-repair] Pass {pass_num}: removing {len(orphaned)} "
                            f"orphaned toolResult blocks at index {i}"
                        )
                        cleaned_blocks = [
                            b for b in content
                            if not (isinstance(b, dict) and "toolResult" in b
                                    and b["toolResult"].get("toolUseId") in orphaned)
                        ]
                        changed = True
                        if cleaned_blocks:
                            new_repaired.append({"role": role, "content": cleaned_blocks})
                        # else: drop empty message entirely
                        continue

                new_repaired.append(msg)

            # --- Assistant messages: validate toolUse against following user ---
            elif role == "assistant":
                tool_use_ids = {
                    b["toolUse"]["toolUseId"]
                    for b in content
                    if isinstance(b, dict) and "toolUse" in b and b["toolUse"].get("toolUseId")
                }
                if tool_use_ids:
                    # Look ahead in *current* (not yet processed) for the next user
                    following_tool_result_ids = set()
                    if i + 1 < len(current) and current[i + 1] is not None:
                        next_msg = current[i + 1]
                        if next_msg.get("role") == "user":
                            next_content = next_msg.get("content", [])
                            if isinstance(next_content, list):
                                following_tool_result_ids = {
                                    b["toolResult"]["toolUseId"]
                                    for b in next_content
                                    if isinstance(b, dict) and "toolResult" in b
                                    and b["toolResult"].get("toolUseId")
                                }

                    orphaned = tool_use_ids - following_tool_result_ids
                    if orphaned:
                        logger.warning(
                            f"[history-repair] Pass {pass_num}: removing {len(orphaned)} "
                            f"orphaned toolUse blocks at index {i}"
                        )
                        cleaned_blocks = [
                            b for b in content
                            if not (isinstance(b, dict) and "toolUse" in b
                                    and b["toolUse"].get("toolUseId") in orphaned)
                        ]
                        changed = True
                        if cleaned_blocks:
                            new_repaired.append({"role": role, "content": cleaned_blocks})
                        # else: drop empty message entirely
                        continue

                new_repaired.append(msg)
            else:
                new_repaired.append(msg)

        current = new_repaired

        if not changed:
            break  # Stable — no more orphans to clean

    if pass_num > 0 and changed:
        logger.info(f"[history-repair] Completed {pass_num + 1} passes to stabilize tool pairs")

    return current


def _prune_conversation_history(history: list, max_messages: int = MAX_HISTORY_MESSAGES) -> list:
    """Prune conversation history to fit within size limits.

    For older messages (beyond the most recent 20), summarize tool blocks
    to reduce size while preserving the fact that tools were called.

    IMPORTANT: The cut boundary respects toolUse/toolResult pairing to
    prevent Bedrock ConverseStream validation errors.
    """
    if len(history) <= max_messages:
        return _validate_tool_pairs(history)

    # Find a safe cut boundary that doesn't orphan toolResult blocks
    desired_start = len(history) - max_messages
    safe_start = _find_safe_cut_index(history, desired_start)
    pruned = history[safe_start:]

    # For older half of messages, compress tool blocks
    compress_boundary = len(pruned) // 2
    for i in range(compress_boundary):
        msg = pruned[i]
        content = msg.get("content")
        if not isinstance(content, list):
            continue
        compressed_blocks = []
        for block in content:
            if not isinstance(block, dict):
                continue
            if "text" in block:
                # Keep text but truncate aggressively for old messages
                text = block["text"]
                if len(text) > 500:
                    text = text[:500] + "... [truncated]"
                compressed_blocks.append({"text": text})
            elif "toolUse" in block:
                tu = block["toolUse"]
                # Keep only tool name and ID, drop input
                compressed_blocks.append({"toolUse": {
                    "name": tu.get("name", ""),
                    "toolUseId": tu.get("toolUseId", ""),
                    "input": {},
                }})
            elif "toolResult" in block:
                tr = block["toolResult"]
                # Keep only status and ID, drop content
                compressed_blocks.append({"toolResult": {
                    "toolUseId": tr.get("toolUseId", ""),
                    "status": tr.get("status", "success"),
                    "content": [{"text": "(previous result)"}],
                }})
        if compressed_blocks:
            pruned[i] = {"role": msg["role"], "content": compressed_blocks}

    # Final validation: ensure no orphaned pairs slipped through
    return _validate_tool_pairs(pruned)


def _sanitize_messages_for_agent(messages: list) -> list:
    """Layer 3 defense: final sanitization before passing messages to Strands Agent.

    Applied after building strands_messages from conversation history, right before
    creating the streaming Agent.  Fixes:
    1. First message must be user role
    2. Consecutive same-role messages merged (role alternation)
    3. toolUse/toolResult pair validation
    4. Trailing assistant orphaned toolUse removal
    5. toolResult count <= preceding toolUse count (the core Bedrock error condition)
    """
    if not messages:
        return messages

    sanitized = []

    # 0. Filter out system-role messages (Bedrock only accepts user/assistant)
    messages = [m for m in messages if m.get("role") in ("user", "assistant")]
    if not messages:
        return []

    # 1. Drop leading non-user messages
    start = 0
    while start < len(messages) and messages[start].get("role") != "user":
        logger.warning(f"[sanitize] Dropping leading {messages[start].get('role')} message")
        start += 1
    if start >= len(messages):
        return []

    # 2. Merge consecutive same-role messages (enforce alternation)
    sanitized = [messages[start]]
    for msg in messages[start + 1:]:
        role = msg.get("role")
        content = msg.get("content")
        if not role:
            continue

        if role == sanitized[-1].get("role"):
            # Merge content blocks
            prev_content = sanitized[-1].get("content", [])
            curr_content = content if isinstance(content, list) else [{"text": str(content)}]
            if isinstance(prev_content, list):
                sanitized[-1] = {"role": role, "content": prev_content + curr_content}
            else:
                # prev was string, curr may be list
                merged_content = [{"text": str(prev_content)}] + curr_content
                sanitized[-1] = {"role": role, "content": merged_content}
            logger.warning(f"[sanitize] Merged consecutive {role} messages")
        else:
            sanitized.append(msg)

    # 3. Validate tool pairs
    sanitized = _validate_tool_pairs(sanitized)

    # 4. Remove orphaned toolUse from last assistant (no following toolResult)
    if sanitized and sanitized[-1].get("role") == "assistant":
        last_content = sanitized[-1].get("content", [])
        if isinstance(last_content, list):
            has_tool_use = any(isinstance(b, dict) and "toolUse" in b for b in last_content)
            if has_tool_use:
                cleaned = [b for b in last_content if not (isinstance(b, dict) and "toolUse" in b)]
                if cleaned:
                    sanitized[-1] = {"role": "assistant", "content": cleaned}
                else:
                    sanitized.pop()
                logger.warning("[sanitize] Removed trailing orphaned toolUse blocks")

    # 5. Final check: toolResult count <= preceding toolUse count per message
    for i in range(len(sanitized)):
        msg = sanitized[i]
        if msg.get("role") != "user":
            continue
        content = msg.get("content", [])
        if not isinstance(content, list):
            continue

        tool_result_ids = [
            b["toolResult"]["toolUseId"]
            for b in content
            if isinstance(b, dict) and "toolResult" in b and b["toolResult"].get("toolUseId")
        ]
        if not tool_result_ids:
            continue

        # Get preceding assistant toolUse IDs
        preceding_tool_use_ids = set()
        if i > 0 and sanitized[i - 1].get("role") == "assistant":
            prev_content = sanitized[i - 1].get("content", [])
            if isinstance(prev_content, list):
                preceding_tool_use_ids = {
                    b["toolUse"]["toolUseId"]
                    for b in prev_content
                    if isinstance(b, dict) and "toolUse" in b and b["toolUse"].get("toolUseId")
                }

        excess_ids = set(tool_result_ids) - preceding_tool_use_ids
        if excess_ids:
            logger.warning(
                f"[sanitize] toolResult count ({len(tool_result_ids)}) > toolUse count "
                f"({len(preceding_tool_use_ids)}) at msg {i}; removing {len(excess_ids)} excess"
            )
            cleaned_blocks = [
                b for b in content
                if not (isinstance(b, dict) and "toolResult" in b
                        and b["toolResult"].get("toolUseId") in excess_ids)
            ]
            if cleaned_blocks:
                sanitized[i] = {"role": "user", "content": cleaned_blocks}
            else:
                sanitized[i] = {"role": "user", "content": [{"text": "(tool results removed)"}]}

    # Final: remove any empty messages that may have been created
    sanitized = [m for m in sanitized if m and m.get("content")]

    return sanitized


# ========================================
# Health Check
# ========================================
@app.get("/ping")
async def ping():
    # NFS mount diagnostics
    s3files_mount = os.environ.get("S3FILES_MOUNT_PATH", "/mnt/s3")
    mount_exists = os.path.isdir(s3files_mount)
    sessions_dir = os.path.join(s3files_mount, "sessions")
    sessions_exists = os.path.isdir(sessions_dir)
    session_count = 0
    if sessions_exists:
        try:
            session_count = len(os.listdir(sessions_dir))
        except OSError:
            session_count = -1

    return JSONResponse(
        content={
            "status": "healthy",
            "mode": "ecs",
            "active_sessions": len(session_store),
            "active_ws": len(_active_ws_connections),
            "timestamp": datetime.utcnow().isoformat(),
            "nfs": {
                "mount_path": s3files_mount,
                "mount_exists": mount_exists,
                "sessions_dir_exists": sessions_exists,
                "session_dirs_count": session_count,
            },
        },
        status_code=200,
    )

# ========================================
# Workspace REST API (for File Explorer)
# ========================================

# Language detection from file extension
_LANG_MAP = {
    ".py": "python", ".js": "javascript", ".ts": "typescript", ".tsx": "typescript",
    ".yaml": "yaml", ".yml": "yaml", ".json": "json", ".md": "markdown",
    ".html": "html", ".css": "css", ".sh": "bash", ".sql": "sql",
    ".xml": "xml", ".txt": "text", ".tf": "hcl", ".java": "java",
    ".csv": "csv", ".tsv": "csv",
    ".docx": "docx", ".pdf": "pdf", ".xlsx": "xlsx",
    ".rb": "ruby", ".go": "go", ".rs": "rust", ".c": "c", ".cpp": "cpp",
    ".h": "c", ".hpp": "cpp", ".r": "r", ".toml": "toml", ".ini": "ini",
    ".env": "text", ".cfg": "text", ".conf": "text", ".log": "text",
}

# Extensions that require binary parsing (not readable as UTF-8 text)
_BINARY_PARSERS = {".docx", ".pdf", ".xlsx"}


def _detect_language(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    return _LANG_MAP.get(ext, "text")


def _read_binary_file(file_path: str, ext: str) -> dict:
    """Extract text content from binary document formats."""
    try:
        if ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table contents
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))
            content = "\n".join(paragraphs)
            return {"success": True, "content": content, "size": len(content)}

        elif ext == ".pdf":
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
            content = "\n\n".join(pages) if pages else "(No extractable text in PDF)"
            return {"success": True, "content": content, "size": len(content)}

        elif ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            sheets_csv = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    rows.append(",".join(f'"{v}"' if "," in v or '"' in v else v for v in cells))
                if rows:
                    header = f"# Sheet: {sheet_name}" if len(wb.sheetnames) > 1 else ""
                    csv_text = "\n".join(rows)
                    sheets_csv.append(f"{header}\n{csv_text}" if header else csv_text)
            wb.close()
            content = "\n\n".join(sheets_csv) if sheets_csv else "(Empty spreadsheet)"
            return {"success": True, "content": content, "size": len(content)}

        return {"success": False, "error": f"Unsupported binary format: {ext}"}
    except ImportError as e:
        logger.warning(f"[workspace-api] Missing parser library for {ext}: {e}")
        return {"success": False, "error": f"Parser not available for {ext} files"}
    except Exception as e:
        logger.error(f"[workspace-api] Failed to parse {ext} file: {e}")
        return {"success": False, "error": f"Failed to parse {ext} file: {str(e)}"}


async def verify_token(authorization: str = Header(default="")):
    """Reusable dependency for JWT validation on REST endpoints."""
    if not os.environ.get("USER_POOL_ID"):
        return {"sub": "local-dev"}
    token = authorization.replace("Bearer ", "") if authorization.startswith("Bearer ") else authorization
    claims = await validate_cognito_token(token)
    if claims is None:
        raise HTTPException(status_code=401, detail="Unauthorized")
    return claims


def _build_tree(directory: str, current_depth: int = 0, max_depth: int = 3) -> list:
    """Recursively build file tree from directory."""
    entries = []
    try:
        with os.scandir(directory) as it:
            for entry in sorted(it, key=lambda e: (not e.is_dir(), e.name)):
                node: Dict[str, Any] = {"name": entry.name, "type": "dir" if entry.is_dir() else "file"}
                if entry.is_file():
                    try:
                        node["size"] = entry.stat().st_size
                    except OSError:
                        node["size"] = 0
                if entry.is_dir() and current_depth < max_depth:
                    node["children"] = _build_tree(entry.path, current_depth + 1, max_depth)
                elif entry.is_dir():
                    node["children"] = []
                entries.append(node)
    except (OSError, PermissionError) as e:
        logger.warning(f"[workspace-api] scandir failed: {e}")
    return entries


@app.get("/api/debug/nfs")
async def debug_nfs():
    """NFS mount diagnostics — no auth required for quick debugging."""
    s3files_mount = os.environ.get("S3FILES_MOUNT_PATH", "/mnt/s3")
    mount_exists = os.path.isdir(s3files_mount)
    sessions_dir = os.path.join(s3files_mount, "sessions")
    sessions_exists = os.path.isdir(sessions_dir)
    session_count = 0
    session_ids: list = []
    if sessions_exists:
        try:
            dirs = os.listdir(sessions_dir)
            session_count = len(dirs)
            # Show last 10 session dirs (sorted by modification time, newest first)
            full_paths = [(d, os.path.getmtime(os.path.join(sessions_dir, d))) for d in dirs]
            full_paths.sort(key=lambda x: x[1], reverse=True)
            session_ids = [d for d, _ in full_paths[:10]]
        except OSError as e:
            session_count = -1
            session_ids = [f"error: {e}"]

    # Check mount contents at root level
    mount_contents: list = []
    if mount_exists:
        try:
            mount_contents = os.listdir(s3files_mount)[:20]
        except OSError:
            mount_contents = ["error listing"]

    return JSONResponse(content={
        "mount_path": s3files_mount,
        "mount_exists": mount_exists,
        "mount_contents": mount_contents,
        "sessions_dir_exists": sessions_exists,
        "session_dirs_count": session_count,
        "recent_sessions": session_ids,
    })


@app.get("/api/workspace/{session_id}/tree")
async def get_workspace_tree(
    session_id: str,
    path: str = "",
    depth: int = 6,
    _claims: dict = Depends(verify_token),
):
    """Return recursive file tree for the session workspace."""
    from tools.workspace_file_tools import _get_session_root, _resolve_safe_path, _S3FILES_MOUNT
    from pathlib import Path as _Path

    try:
        if path:
            target = _resolve_safe_path(session_id, path)
        else:
            target = _get_session_root(session_id)

        mount_path = _Path(_S3FILES_MOUNT)
        mount_exists = mount_path.exists()
        sessions_dir = mount_path / "sessions"
        sessions_exists = sessions_dir.exists()

        if not target.exists():
            # Log diagnostic info for debugging NFS mount issues
            mount_contents = "N/A"
            if mount_exists:
                try:
                    mount_contents = str([e.name for e in mount_path.iterdir()])
                except Exception:
                    mount_contents = "error listing"
            logger.warning(
                f"[workspace-api] tree: session dir not found. "
                f"target={target}, mount_exists={mount_exists}, "
                f"sessions_exists={sessions_exists}, "
                f"mount_contents={mount_contents}"
            )
            return JSONResponse({"success": True, "tree": []})

        tree = _build_tree(str(target), max_depth=min(depth, 8))
        return JSONResponse({"success": True, "tree": tree})
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"[workspace-api] tree error: {e}")
        raise HTTPException(status_code=500, detail="Internal error")


@app.get("/api/workspace/{session_id}/file")
async def get_workspace_file(
    session_id: str,
    path: str = Query(...),
    _claims: dict = Depends(verify_token),
):
    """Return file content from the session workspace."""
    from tools.workspace_file_tools import read_workspace_file as _read_file, _resolve_safe_path

    try:
        ext = os.path.splitext(path)[1].lower()

        # For known binary formats, skip text read and go straight to parser
        if ext in _BINARY_PARSERS:
            resolved = str(_resolve_safe_path(session_id, path))
            result = _read_binary_file(resolved, ext)
            if not result.get("success"):
                raise HTTPException(status_code=422, detail=result.get("error", "Cannot parse file"))
            # xlsx → return as csv language for table rendering
            lang = "csv" if ext == ".xlsx" else _detect_language(path)
            return JSONResponse({
                "success": True,
                "content": result["content"],
                "size": result["size"],
                "language": lang,
            })

        # Standard text read
        result = _read_file(session_id=session_id, path=path)
        if not result.get("success"):
            raise HTTPException(status_code=404, detail=result.get("error", "File not found"))
        return JSONResponse({
            "success": True,
            "content": result["content"],
            "size": result["size"],
            "language": _detect_language(path),
        })
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"[workspace-api] file error: {e}")
        raise HTTPException(status_code=500, detail="Internal error")


# ========================================
# Asset Download REST API (always fresh from S3)
# ========================================
@app.get("/api/assets/{session_id}/download")
async def download_assets(
    session_id: str,
    asset_type: Optional[str] = Query(default=None, description="Filter by asset type; omit or 'all' for full package"),
    _claims: dict = Depends(verify_token),
):
    """Package session assets directly from S3 (bypasses NFS cache) and return a presigned URL.

    Called on every click so the download always reflects the latest S3 content,
    regardless of NFS cache or cached packageS3Key in the frontend.
    """
    from tools.asset_packager import package_assets_impl

    filter_value = None if not asset_type or asset_type.lower() == "all" else asset_type
    result = package_assets_impl(
        session_id=session_id,
        asset_type_filter=filter_value,
        include_readme=filter_value is None,
    )
    if not result.get("success"):
        raise HTTPException(status_code=404, detail=result.get("error", "No assets found"))
    return JSONResponse({
        "success": True,
        "downloadUrl": result["download_url"],
        "expiresAt": result["expires_at"],
        "s3Key": result["s3_key"],
        "fileCount": result["file_count"],
    })


# ========================================
# Message Log REST API (for catch-up after reconnect)
# ========================================
@app.get("/api/message-log/{session_id}")
async def get_message_log_endpoint(
    session_id: str,
    after_seq: int = Query(default=0),
    _claims: dict = Depends(verify_token),
):
    """Return message log entries after the given sequence number.

    Used by the frontend to catch up on events missed while disconnected.
    """
    try:
        msg_log = get_message_log(S3FILES_MOUNT, session_id)
        entries = msg_log.read_after(after_seq)

        # Check if a background task is still running for this session
        is_active = False
        async with _background_tasks_lock:
            bg = _background_tasks.get(session_id)
            if bg and not bg["task"].done():
                is_active = True

        return JSONResponse({
            "entries": entries,
            "isAgentActive": is_active,
        })
    except Exception as e:
        logger.error(f"[message-log] Error reading log for {session_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to read message log")


# ========================================
# HTTP Invocations
# ========================================
class InvocationRequest(BaseModel):
    action: str = "chat"
    prompt: Optional[str] = None
    session_id: str = "default"
    document_content: Optional[str] = None

@app.post("/invocations")
async def invoke(request: InvocationRequest):
    session_id = request.session_id
    action = request.action

    try:
        if action == "chat":
            if not request.prompt:
                return JSONResponse({"success": False, "error": "No prompt provided"})
            session = get_or_create_session(session_id)
            agent = session["agent"]
            session["conversation_history"].append({"role": "user", "content": request.prompt})
            result = agent(request.prompt)
            response_text = str(result)
            summarized = summarize_response_for_history(response_text)
            session["conversation_history"].append({"role": "assistant", "content": summarized})
            return JSONResponse({"success": True, "response": response_text, "session_id": session_id})
        elif action == "get_progress":
            if session_id not in session_store:
                return JSONResponse({"success": True, "progress": {"message_count": 0}})
            session = session_store[session_id]
            return JSONResponse({
                "success": True,
                "progress": {
                    "message_count": len(session.get("conversation_history", [])),
                    "document_mode": session.get("document_mode", False),
                },
            })
        else:
            return JSONResponse({"success": False, "error": f"Unknown action: {action}"})
    except Exception as e:
        logger.error(f"Invocation error: {e}\n{traceback.format_exc()}")
        return JSONResponse({"success": False, "error": str(e)})

# ========================================
# WebSocket Safe Send
# ========================================
WS_PAYLOAD_WARN_BYTES = 15_000
WS_PAYLOAD_MAX_BYTES = 32_000

async def safe_send_json(websocket: WebSocket, data: dict) -> bool:
    """Send JSON over WebSocket with payload size defense."""
    try:
        payload = json.dumps(data, ensure_ascii=False)
        payload_len = len(payload.encode("utf-8"))

        if payload_len > WS_PAYLOAD_MAX_BYTES:
            # Truncate content fields
            if "content" in data and isinstance(data["content"], str):
                original_len = len(data["content"])
                excess_ratio = WS_PAYLOAD_MAX_BYTES / payload_len
                trim_target = int(original_len * excess_ratio * 0.9)
                data["content"] = data["content"][:trim_target] + "\n... [truncated]"
            elif "input" in data and isinstance(data["input"], dict):
                for key, val in data["input"].items():
                    if isinstance(val, str) and len(val) > 500:
                        data["input"][key] = val[:200] + f"\n... [{len(val)} chars, truncated]"
            payload = json.dumps(data, ensure_ascii=False)

        await websocket.send_json(data)
        return True
    except WebSocketDisconnect:
        return False
    except Exception:
        return False

# ========================================
# Tool Result Helpers
# ========================================

def _extract_tool_result_text(content_blocks: list) -> str:
    """Extract readable text from Strands ToolResultContent blocks."""
    if not isinstance(content_blocks, list):
        return str(content_blocks) if content_blocks else "(no result)"

    texts = []
    for block in content_blocks:
        if isinstance(block, dict):
            if "text" in block:
                texts.append(block["text"])
            elif "json" in block:
                texts.append(json.dumps(block["json"], ensure_ascii=False, default=str))

    result = "\n".join(texts) if texts else str(content_blocks)

    # Strands wraps dict returns as str(dict) — try to make it more readable
    if result and (result.startswith("{") or result.startswith("\\u007b") or result.startswith("{'") or result.startswith('{"')):
        try:
            import ast
            parsed = ast.literal_eval(result)
            if isinstance(parsed, dict):
                result = _format_tool_result_summary(parsed)
        except (ValueError, SyntaxError):
            pass

    # Truncate
    if len(result) > 1500:
        result = result[:1500] + "..."
    return result


def _format_tool_result_summary(result_dict: dict) -> str:
    """Format tool result dict into a concise, readable summary."""
    parts = []
    if "message" in result_dict:
        parts.append(str(result_dict["message"]))
    if "summary" in result_dict:
        s = result_dict["summary"]
        if isinstance(s, str):
            parts.append(s)
        elif isinstance(s, dict):
            for k, v in list(s.items())[:5]:
                parts.append(f"{k}: {v}")
    if "count" in result_dict:
        parts.append(f"count: {result_dict['count']}")
    if "operation_id" in result_dict:
        parts.append(f"operation: {result_dict['operation_id']}")
    if "text_summary" in result_dict:
        t = str(result_dict["text_summary"])
        parts.append(t[:200] if len(t) > 200 else t)

    if parts:
        return " | ".join(parts[:3])

    return json.dumps(result_dict, ensure_ascii=False, default=str)[:500]


# ========================================
# WebSocket Endpoint
# ========================================
@app.websocket("/ws")
async def websocket_handler(
    websocket: WebSocket,
    token: str = Query(default=""),
    sessionId: str = Query(default=""),
):
    """
    WebSocket endpoint for real-time bidirectional streaming.

    Connection: wss://{ALB_DNS}/ws?token={idToken}&sessionId={session_id}
    """
    # Validate Cognito token
    claims = await validate_cognito_token(token)
    if claims is None and os.environ.get("USER_POOL_ID"):
        await websocket.close(code=4001, reason="Unauthorized")
        return

    await websocket.accept()

    session_id = sessionId or f"ws-{id(websocket)}"
    _active_ws_connections.add(session_id)
    logger.info(f"[WS] Connected: {session_id} (active: {len(_active_ws_connections)})")

    # Send connected event (include current phase + NFS progress for frontend restoration)
    _conn_progress = _get_frontend_progress(session_id)
    await safe_send_json(websocket, {
        "type": "connected",
        "sessionId": session_id,
        "phase": _detect_phase(session_id),
        "progressState": _conn_progress if _conn_progress else None,
    })

    # Check for running background task and reattach WebSocket
    async with _background_tasks_lock:
        bg = _background_tasks.get(session_id)
        if bg and not bg["task"].done():
            logger.info(f"[WS] Reattaching WebSocket to running background task for {session_id}")
            bg["ws_holder"]["ws"] = websocket
            await safe_send_json(websocket, {
                "type": "background_task_active",
                "sessionId": session_id,
                "message": "Agent is still processing your request...",
            })

    client_connected = True

    try:
        while True:
            data = await websocket.receive_json()
            action = data.get("action", "sendMessage")

            # Bind the current ContextVar session_id for this WS iteration so
            # every downstream module (spec_manager, streaming_callback, agents,
            # etc.) reads the correct per-session state even before action
            # handlers run set_workspace_session_id().  This is the core of the
            # session-isolation fix for the ECS (single-process) runtime.
            try:
                _sess = session_store.get(session_id, {})
                _eff_sid = (
                    _sess.get("session_data", {}).get("original_session_id", session_id)
                    if isinstance(_sess, dict) else session_id
                )
                from tools.session_context import current_session_id as _current_sid
                _current_sid.set(_eff_sid)
            except Exception as _bind_err:
                logger.warning(f"[WS] session_id bind failed: {_bind_err}")

            if action == "sendMessage":
                await handle_send_message_ws(websocket, session_id, data)
            elif action == "sendMessageWithAttachments":
                await handle_send_message_with_attachments_ws(websocket, session_id, data)
            elif action == "sendMessageWithS3Attachments":
                await handle_send_message_with_s3_attachments_ws(websocket, session_id, data)
            elif action == "uploadQuestionnaire":
                await handle_upload_questionnaire_ws(websocket, session_id, data)
            elif action in ("getAssets", "downloadAssets"):
                await handle_get_assets_ws(websocket, session_id)
            elif action == "getProgress":
                await handle_get_progress_ws(websocket, session_id)
            elif action == "injectHistory":
                await handle_inject_history_ws(websocket, session_id, data)
            elif action == "createNewSession":
                await handle_create_new_session_ws(websocket, session_id)
            elif action == "ping":
                await safe_send_json(websocket, {"type": "pong"})
            else:
                await safe_send_json(websocket, {"type": "error", "content": f"Unknown action: {action}"})

    except WebSocketDisconnect as e:
        client_connected = False
        logger.info(f"[WS] Disconnected: {session_id} (code: {getattr(e, 'code', 'unknown')})")
    except Exception as e:
        logger.error(f"[WS] Error: {session_id}: {e}\n{traceback.format_exc()}")
        if client_connected:
            await safe_send_json(websocket, {
                "type": "error",
                "content": str(e),
                "debug": {"error_type": type(e).__name__, "session_id": session_id},
            })
        try:
            await websocket.close()
        except Exception:
            pass
    finally:
        _active_ws_connections.discard(session_id)

        # If a background task is still running, detach the WebSocket but let it continue
        async with _background_tasks_lock:
            bg = _background_tasks.get(session_id)
            if bg and not bg["task"].done():
                logger.info(f"[WS] Detaching WebSocket from running background task for {session_id}")
                bg["ws_holder"]["ws"] = None
                # History will be saved when the background task completes
            else:
                # No background task — persist conversation history to NFS on disconnect
                if session_id in session_store:
                    try:
                        history = session_store[session_id].get("conversation_history", [])
                        if history:
                            _context_store.save_conversation_history(session_id, history)
                    except Exception as e:
                        logger.warning(f"History flush failed on disconnect for {session_id}: {e}")
        logger.info(f"[WS] Cleanup: {session_id} (active: {len(_active_ws_connections)})")

# ========================================
# WebSocket Action Handlers
# ========================================

async def handle_send_message_ws(websocket: WebSocket, session_id: str, message: Dict[str, Any]):
    """Handle chat message with streaming response.

    The agent runs as a background asyncio.Task so it survives WebSocket
    disconnects.  All events are logged to NFS (MessageLog) and best-effort
    forwarded to the currently attached WebSocket.
    """
    # Guard: reject if an agent task is already running for this session
    async with _background_tasks_lock:
        bg = _background_tasks.get(session_id)
        if bg and not bg["task"].done():
            await safe_send_json(websocket, {
                "type": "error",
                "content": "Agent is still processing your previous request. Please wait.",
            })
            return

    user_message = message.get("message", "")
    content_blocks = message.get("content_blocks")  # Multimodal content from attachment handlers
    if not user_message and not content_blocks:
        await safe_send_json(websocket, {"type": "error", "content": "No message provided"})
        return

    session = get_or_create_session(session_id)

    # Send typing indicator
    if not await safe_send_json(websocket, {"type": "typing", "status": "Agent is thinking..."}):
        return

    # Set session ID for S3 asset storage
    effective_session_id = session.get("session_data", {}).get("original_session_id", session_id)
    set_streaming_session_id(effective_session_id)

    # Initialise S3 project workspace
    try:
        from tools.project_workspace import set_workspace_session_id
        set_workspace_session_id(effective_session_id)
    except Exception as e:
        logger.warning(f"workspace_init_failed: {e}")

    set_message_index(len(session["conversation_history"]))

    # Set up Sub-Agent callback handler for streaming
    from agents.streaming_handler import StrandsCallbackHandler as SubAgentCallbackHandler
    try:
        callback_handler = SubAgentCallbackHandler(None, websocket)
        callback_handler.set_event_loop(asyncio.get_running_loop())
        callback_handler.setup_streaming_callback()
        # Main stream_async loop sends tool_start/tool_end/thinking directly.
        # Suppress these from callback handler to avoid duplication and latency.
        callback_handler.suppress_tool_ws_events = True
    except Exception:
        callback_handler = None

    if callback_handler:
        set_research_callback(callback_handler)
        set_faq_callback(callback_handler)
        set_lambda_callback(callback_handler)
        set_openapi_callback(callback_handler)
        set_prompt_callback(callback_handler)
        set_contact_flow_callback(callback_handler)
        set_infrastructure_callback(callback_handler)
        set_reviewer_callback(callback_handler)

    # Build Strands messages from history.
    # History can be in two formats:
    # 1. Legacy text-only: {"role": "user", "content": "text"}
    # 2. Full Strands format: {"role": "user", "content": [{"text": "..."}, {"toolUse": ...}]}
    strands_messages = []
    for msg in session["conversation_history"]:
        content = msg.get("content")
        if isinstance(content, str):
            # Legacy text-only format → convert to Strands content block
            strands_messages.append({
                "role": msg["role"],
                "content": [{"text": content}],
            })
        elif isinstance(content, list):
            # Already in Strands format (has tool_use/tool_result blocks)
            strands_messages.append(msg)
        else:
            # Skip malformed entries
            continue

    # Layer 3: sanitize messages before passing to agent
    strands_messages = _sanitize_messages_for_agent(strands_messages)

    # Capture message count before stream_async appends new messages
    pre_stream_message_count = len(strands_messages)

    # Dynamic phase-based system prompt selection
    current_phase = _detect_phase(effective_session_id)

    # Context boundary: interview → generation handoff
    if current_phase != "interview" and not session.get("_handoff_processed"):
        handoff = check_interview_handoff(effective_session_id)
        if handoff:
            logger.info(f"[handoff] Interview→Generation context boundary for {effective_session_id}")
            # Archive interview history to NFS before clearing
            try:
                import json as _json
                from pathlib import Path as _Path
                _mount_root = _Path(S3FILES_MOUNT).resolve()
                _safe_session = _sanitize_path_component(effective_session_id)
                archive_path = (_mount_root / "sessions" / _safe_session / "context" / "interview_history.json").resolve()
                if _mount_root not in archive_path.parents:
                    raise ValueError(f"archive path escapes mount root: {archive_path}")
                archive_path.parent.mkdir(parents=True, exist_ok=True)
                with open(archive_path, "w", encoding="utf-8") as _f:
                    _json.dump(session["conversation_history"], _f, ensure_ascii=False, default=str)
            except Exception as arch_err:
                logger.warning(f"[handoff] history archive failed (non-critical): {arch_err}")
            # Clear context for fresh generation start
            session["conversation_history"] = []
            strands_messages = []
            session["_handoff_processed"] = True
            # Persist phase transition to NFS (prevents duplicate phase_changed on first sub-agent completion)
            _update_phase(effective_session_id, "generation", "interview_complete_handoff")
            # Notify frontend
            await safe_send_json(websocket, {"type": "phase_changed", "phase": "generation", "previousPhase": "interview", "context_cleared": True})
            # Inject bootstrap message so generation agent knows where specs are
            bootstrap_msg = (
                "[System] Interview phase complete. All specifications have been saved to the workspace.\n"
                "Load specifications using:\n"
                "- get_all_tool_ids() to enumerate all operations and tools\n"
                "- get_operation_spec(operation_id) for individual operation details\n"
                "- get_session_flow_config_tool() for session-level config\n"
                "- load_requirement_document(doc_type=\"analysis\") for the requirements analysis document\n"
                "Begin generation following the 5-phase workflow."
            )
            strands_messages = [{"role": "user", "content": [{"text": bootstrap_msg}]}]

    phase_prompt = get_phase_system_prompt(current_phase)
    phase_tools = get_tools_for_phase(current_phase)
    logger.info(f"[phase] Using phase '{current_phase}' system prompt for {effective_session_id}")

    # Create streaming agent
    streaming_agent = Agent(
        model=session["model"],
        system_prompt=phase_prompt,
        tools=phase_tools,
        callback_handler=callback_handler,
        messages=strands_messages,
    )

    # Session context prefix
    ui_language = message.get("language", "ko-KR")
    context_prefix = f'[Session: session_id="{effective_session_id}" language="{ui_language}"]'

    # Structured Note-Taking: inject generation progress into context
    # This survives conversation history pruning and gives the orchestrator
    # full awareness of what has been generated/reviewed/fixed.
    generation_state_block = ""
    try:
        progress = _read_generation_progress(effective_session_id)
        if progress:
            generation_state_block = (
                "\n<generation_state>\n"
                "Below is the authoritative log of all generation, review, and fix events.\n"
                "Trust this over conversation memory. Do NOT regenerate assets marked ✅.\n"
                "Assets marked 🔧→✅ are already fixed — do not re-fix.\n"
                "⛔ RULE: Only fix assets that the reviewer flagged AND the user EXPLICITLY confirmed.\n"
                "⛔ RULE: Complete ONE phase per turn, then STOP and wait for user.\n\n"
                f"{progress}\n"
                "</generation_state>\n\n"
            )
    except Exception as prog_err:
        logger.warning(f"[generation_progress] read failed (non-critical): {prog_err}")

    # Modification tracking: record this turn's keywords & bump counters,
    # then inject a <modification_state> block so the orchestrator can route
    # edits to the right asset and detect repeated corrections.
    modification_state_block = ""
    try:
        raw_user_text = user_message if isinstance(user_message, str) else ""
        if raw_user_text:
            record_modification_request(effective_session_id, raw_user_text)
        mod_state = format_modification_state_block(effective_session_id, raw_user_text)
        if mod_state:
            modification_state_block = (
                "\n<modification_state>\n"
                f"{mod_state}\n"
                "</modification_state>\n\n"
            )
    except Exception as mod_err:
        logger.warning(f"[modification_tracking] inject failed (non-critical): {mod_err}")

    combined_state = f"{generation_state_block}{modification_state_block}"

    if content_blocks:
        # Multimodal: prepend session context to the text block in content_blocks
        text_found = False
        for i, block in enumerate(content_blocks):
            if "text" in block:
                content_blocks[i] = {"text": f"{context_prefix}{combined_state}\n\n{block['text']}"}
                text_found = True
                break
        if not text_found:
            content_blocks.append({"text": f"{context_prefix}{combined_state}\n\nPlease analyze the attached files."})
        message_for_agent = content_blocks
    else:
        message_for_agent = f'{context_prefix}{combined_state}\n\n{user_message}'

    # ── NFS Message Log + Background Task ──
    msg_log = get_message_log(S3FILES_MOUNT, session_id)
    msg_log.clear()  # fresh log per agent invocation

    # Mutable WebSocket reference — set to None on disconnect, replaced on reconnect
    ws_holder: Dict[str, Any] = {"ws": websocket}

    async def safe_send_or_log(data: dict) -> bool:
        """Always log to NFS, best-effort send to WebSocket."""
        msg_log.append(data)
        ws = ws_holder.get("ws")
        if ws is not None:
            try:
                return await safe_send_json(ws, data)
            except Exception:
                return False
        return False  # no WebSocket attached

    # ── Background agent task ──
    async def run_agent_background():
        """Run the agent streaming loop. Survives WebSocket disconnects."""
        full_response = ""
        tool_invocations_started: set = set()
        tool_inputs: dict = {}      # toolUseId → latest input dict
        tool_names_map: dict = {}   # toolUseId → tool name

        # Heartbeat loop
        flush_stop = asyncio.Event()

        async def heartbeat_loop():
            seq = 0
            while not flush_stop.is_set():
                try:
                    await asyncio.wait_for(flush_stop.wait(), timeout=2.0)
                    break
                except asyncio.TimeoutError:
                    seq += 1
                    if callback_handler:
                        ws_events = callback_handler.get_pending_ws_events()
                        for evt in ws_events:
                            await safe_send_or_log(evt)
                    else:
                        await safe_send_or_log({"type": "heartbeat", "seq": seq})

        heartbeat_task = asyncio.create_task(heartbeat_loop())

        try:
            async for event in streaming_agent.stream_async(message_for_agent):
                # Thinking / reasoning blocks (Claude extended thinking)
                if "reasoningText" in event:
                    thinking_chunk = event["reasoningText"]
                    if thinking_chunk:
                        await safe_send_or_log({"type": "thinking", "content": thinking_chunk})

                # Text streaming
                elif "data" in event:
                    chunk = event["data"]
                    full_response += chunk
                    await safe_send_or_log({"type": "stream", "content": chunk})

                # Tool use start
                elif "current_tool_use" in event:
                    tool_use = event["current_tool_use"]
                    tool_name = tool_use.get("name", "")
                    tool_use_id = tool_use.get("toolUseId", "")

                    tool_input = tool_use.get("input", {})
                    if tool_use_id and isinstance(tool_input, dict) and tool_input:
                        tool_inputs[tool_use_id] = tool_input

                    if tool_use_id and tool_use_id not in tool_invocations_started:
                        tool_invocations_started.add(tool_use_id)
                        tool_names_map[tool_use_id] = tool_name

                        progress_id = SUBAGENT_TO_PROGRESS_ID.get(tool_name)
                        if progress_id:
                            await safe_send_or_log({
                                "type": "progress_update",
                                "itemId": progress_id,
                                "status": "generating",
                            })

                        await safe_send_or_log({
                            "type": "tool_start",
                            "tool": tool_name,
                            "toolUseId": tool_use_id,
                            "input": tool_input,
                        })

                # ToolResultMessageEvent
                elif "message" in event:
                    msg = event["message"]
                    if isinstance(msg, dict) and msg.get("role") == "user":
                        for block in (msg.get("content") or []):
                            if isinstance(block, dict) and "toolResult" in block:
                                tr = block["toolResult"]
                                tr_tool_use_id = tr.get("toolUseId", "")
                                tr_tool_name = tool_names_map.get(tr_tool_use_id, "unknown")

                                result_text = _extract_tool_result_text(tr.get("content", []))
                                status = tr.get("status", "success")
                                if status == "success":
                                    status = "completed"

                                progress_id = SUBAGENT_TO_PROGRESS_ID.get(tr_tool_name)
                                if progress_id:
                                    await safe_send_or_log({
                                        "type": "progress_update",
                                        "itemId": progress_id,
                                        "status": "completed",
                                    })

                                # Record tool completion to NFS (real-time, inline)
                                try:
                                    _record_tool_completion(
                                        effective_session_id, tr_tool_name, status
                                    )
                                except Exception:
                                    pass  # non-critical

                                # Phase transition detection
                                try:
                                    if progress_id:  # sub-agent completion
                                        old_phase = _read_phase(effective_session_id)
                                        detected = _detect_phase(effective_session_id)
                                        if old_phase == "interview" and detected != "interview":
                                            _update_phase(effective_session_id, "generation", f"first_subagent:{tr_tool_name}")
                                            await safe_send_or_log({"type": "phase_changed", "phase": "generation", "previousPhase": "interview"})
                                        elif old_phase == "generation" and detected == "review":
                                            _update_phase(effective_session_id, "review", "all_core_assets_completed")
                                            await safe_send_or_log({"type": "phase_changed", "phase": "review", "previousPhase": "generation"})
                                        elif old_phase == "review" and detected == "post_generation":
                                            _update_phase(effective_session_id, "post_generation", f"post_review_fix:{tr_tool_name}")
                                            await safe_send_or_log({"type": "phase_changed", "phase": "post_generation", "previousPhase": "review"})
                                    elif tr_tool_name == "save_operation_spec":
                                        new_phase = _read_phase(effective_session_id)
                                        if new_phase != current_phase:
                                            await safe_send_or_log({"type": "phase_changed", "phase": new_phase, "previousPhase": current_phase})
                                except Exception:
                                    pass  # phase tracking is non-critical

                                await safe_send_or_log({
                                    "type": "tool_end",
                                    "tool": tr_tool_name,
                                    "toolUseId": tr_tool_use_id,
                                    "input": tool_inputs.get(tr_tool_use_id, {}),
                                    "result": result_text,
                                    "status": status,
                                })

                                tool_invocations_started.discard(tr_tool_use_id)

            # Send tool_end for any tools that started but never got tool_result
            for remaining_id in list(tool_invocations_started):
                await safe_send_or_log({
                    "type": "tool_end",
                    "tool": tool_names_map.get(remaining_id, "unknown"),
                    "toolUseId": remaining_id,
                    "input": tool_inputs.get(remaining_id, {}),
                    "result": "(completed)",
                    "status": "completed",
                })

            # Flush remaining callback handler events
            if callback_handler:
                remaining_events = callback_handler.get_pending_ws_events()
                for evt in remaining_events:
                    await safe_send_or_log(evt)

            # Stream end
            await safe_send_or_log({"type": "stream_end"})

            # Chat-input placeholder hint — compute once per turn so the
            # frontend textarea can nudge the user with a concrete example
            # appropriate to the current phase / modification state.
            try:
                phase_for_hint = _detect_phase(effective_session_id)
                hint_text = suggest_placeholder(
                    effective_session_id,
                    raw_user_text if isinstance(raw_user_text, str) else "",
                    phase=phase_for_hint,
                    language=ui_language,
                )
                await safe_send_or_log({
                    "type": "input_hint",
                    "placeholder": hint_text,
                    "phase": phase_for_hint,
                })
            except Exception as hint_err:
                logger.warning(f"[input_hint] send failed (non-critical): {hint_err}")

            # Save full Strands messages
            new_messages = _extract_new_messages(streaming_agent.messages, pre_stream_message_count)
            logger.info(
                f"[generation_progress] raw agent messages: {len(streaming_agent.messages)}, "
                f"pre_count: {pre_stream_message_count}, "
                f"new_messages: {len(new_messages)}"
            )

            # Structured Note-Taking: scan sub-agent completions and log progress
            try:
                _update_generation_progress(effective_session_id, new_messages)
            except Exception as prog_err:
                logger.warning(f"[generation_progress] update failed (non-critical): {prog_err}")

            session["conversation_history"].extend(new_messages)
            session["conversation_history"] = _prune_conversation_history(
                session["conversation_history"], max_messages=MAX_HISTORY_MESSAGES
            )
            _context_store.save_conversation_history(session_id, session["conversation_history"])

        except Exception as e:
            logger.error(f"[BG] Streaming error for {session_id}: {e}\n{traceback.format_exc()}")
            await safe_send_or_log({"type": "error", "content": str(e)})
            # Save partial history on error
            try:
                partial = _extract_new_messages(streaming_agent.messages, pre_stream_message_count)
                if partial:
                    # Log progress even on error (captures completions before failure)
                    try:
                        _update_generation_progress(effective_session_id, partial)
                    except Exception:
                        pass
                    session["conversation_history"].extend(partial)
                    session["conversation_history"] = _prune_conversation_history(
                        session["conversation_history"], max_messages=MAX_HISTORY_MESSAGES
                    )
                    _context_store.save_conversation_history(session_id, session["conversation_history"])
            except Exception as save_err:
                logger.error(f"[BG] Failed to save partial history for {session_id}: {save_err}")
        finally:
            flush_stop.set()
            await heartbeat_task
            # Last-resort: ensure conversation history is persisted to NFS
            try:
                history = session.get("conversation_history", [])
                if history:
                    _context_store.save_conversation_history(session_id, history)
            except Exception as final_save_err:
                logger.error(f"[BG] Final history save failed for {session_id}: {final_save_err}")
            # Remove from background tasks registry
            async with _background_tasks_lock:
                _background_tasks.pop(session_id, None)
            logger.info(f"[BG] Agent task finished for {session_id}")

    # Register and launch background task (fire-and-forget).
    # We do NOT await the task here — returning immediately lets the main
    # WebSocket loop continue to read client pings and other actions.
    # Without this, client pings sit unread while the agent executes and the
    # client closes with 1011 "keepalive ping timeout".
    task = asyncio.create_task(run_agent_background())
    async with _background_tasks_lock:
        _background_tasks[session_id] = {
            "task": task,
            "ws_holder": ws_holder,
            "started_at": time.time(),
        }


def _sanitize_path_component(value: str) -> str:
    """Strip path-traversal segments and separators from a single path component."""
    return value.replace("..", "_").replace("/", "_").replace("\\", "_")


def _save_attachments_to_workspace(session_id: str, attachments: list):
    """Save inline (base64) attachments to NFS workspace for FileExplorer visibility."""
    try:
        import base64 as b64
        safe_session = _sanitize_path_component(session_id)
        uploads_dir = os.path.join(S3FILES_MOUNT, "sessions", safe_session, "uploads")
        os.makedirs(uploads_dir, exist_ok=True)
        for att in attachments:
            name = att.get("name", "unknown")
            data_b64 = att.get("data", "")
            if not data_b64:
                continue
            # Strip data URL prefix if present
            if "," in data_b64:
                data_b64 = data_b64.split(",", 1)[1]
            try:
                file_bytes = b64.b64decode(data_b64)
                safe_name = _sanitize_path_component(os.path.basename(name))
                filepath = os.path.join(uploads_dir, safe_name)
                with open(filepath, "wb") as f:
                    f.write(file_bytes)
                logger.info(f"[WORKSPACE] Saved upload: {filepath} ({len(file_bytes)} bytes)")
            except Exception as e:
                logger.warning(f"[WORKSPACE] Failed to save {name}: {e}")
    except Exception as e:
        logger.warning(f"[WORKSPACE] Failed to save attachments: {e}")


def _save_s3_attachments_to_workspace(session_id: str, s3_attachments: list):
    """Save S3 attachments to NFS workspace for FileExplorer visibility."""
    try:
        from tools.attachment_handler import read_file_from_s3
        safe_session = _sanitize_path_component(session_id)
        uploads_dir = os.path.join(S3FILES_MOUNT, "sessions", safe_session, "uploads")
        os.makedirs(uploads_dir, exist_ok=True)
        for att in s3_attachments:
            s3_key = att.get("s3Key", "")
            filename = att.get("filename") or s3_key.split("/")[-1]
            if not s3_key:
                continue
            try:
                file_bytes = read_file_from_s3(s3_key)
                if file_bytes:
                    safe_filename = _sanitize_path_component(os.path.basename(filename))
                    filepath = os.path.join(uploads_dir, safe_filename)
                    with open(filepath, "wb") as f:
                        f.write(file_bytes)
                    logger.info(f"[WORKSPACE] Saved S3 upload: {filepath} ({len(file_bytes)} bytes)")
            except Exception as e:
                logger.warning(f"[WORKSPACE] Failed to save {filename} from S3: {e}")
    except Exception as e:
        logger.warning(f"[WORKSPACE] Failed to save S3 attachments: {e}")


async def handle_send_message_with_attachments_ws(websocket: WebSocket, session_id: str, data: Dict[str, Any]):
    """Handle message with inline attachments (base64) — images and documents."""
    try:
        from tools.attachment_handler import (
            convert_attachments_to_content_blocks,
            validate_attachments,
            format_attachments_for_history,
        )
        attachments = data.get("attachments", [])

        # validate_attachments returns Tuple[bool, str]
        is_valid, error_msg = validate_attachments(attachments)
        if not is_valid:
            await safe_send_json(websocket, {"type": "error", "content": error_msg})
            return

        message_text = data.get("message", "")

        # Convert to Strands multimodal content blocks (images + documents)
        content_blocks = convert_attachments_to_content_blocks(attachments, message_text)
        if content_blocks:
            data["content_blocks"] = content_blocks

        # Save uploads to NFS workspace for FileExplorer visibility
        _save_attachments_to_workspace(session_id, attachments)

        # Store text with attachment markers for conversation history
        markers = format_attachments_for_history(attachments)
        data["message"] = f"{message_text}\n\n{markers}" if message_text else markers

        await handle_send_message_ws(websocket, session_id, data)
    except Exception as e:
        logger.error(f"Attachment handling error: {e}\n{traceback.format_exc()}")
        await handle_send_message_ws(websocket, session_id, data)


async def handle_send_message_with_s3_attachments_ws(websocket: WebSocket, session_id: str, data: Dict[str, Any]):
    """Handle message with S3-based attachments (large files uploaded via presigned URL)."""
    try:
        from tools.attachment_handler import convert_s3_attachments_to_content_blocks
        message_text = data.get("message", "")
        s3_attachments = data.get("s3Attachments", [])

        if s3_attachments:
            # Convert S3 references to Strands multimodal content blocks
            content_blocks = convert_s3_attachments_to_content_blocks(s3_attachments, message_text)
            if content_blocks:
                data["content_blocks"] = content_blocks

            # Save uploads to NFS workspace for FileExplorer visibility
            _save_s3_attachments_to_workspace(session_id, s3_attachments)

            # Store text with attachment markers for conversation history
            markers = " ".join(
                f"[Attached: {att.get('filename', att.get('s3Key', '').split('/')[-1])}]"
                for att in s3_attachments
            )
            data["message"] = f"{message_text}\n\n{markers}" if message_text else markers

        await handle_send_message_ws(websocket, session_id, data)
    except Exception as e:
        logger.error(f"S3 attachment handling error: {e}\n{traceback.format_exc()}")
        await handle_send_message_ws(websocket, session_id, data)


async def handle_upload_questionnaire_ws(websocket: WebSocket, session_id: str, data: Dict[str, Any]):
    """Handle document upload."""
    content = data.get("content", "")
    if not content or not content.strip():
        await safe_send_json(websocket, {"type": "error", "content": "Empty document"})
        return

    session = get_or_create_session(session_id)
    session["uploaded_document"] = content
    session["document_mode"] = True

    document_prompt = get_document_analysis_prompt(content)
    data["message"] = document_prompt + "\n\n문서를 분석하고 고객에게 응답해주세요."
    await handle_send_message_ws(websocket, session_id, data)


async def handle_get_assets_ws(websocket: WebSocket, session_id: str):
    """Package and return generated assets as a downloadable ZIP."""
    try:
        from tools.asset_packager import package_and_upload_assets

        # Derive project name from session context
        session = session_store.get(session_id, {})
        company = session.get("session_data", {}).get("companyName", "")
        project_name = re.sub(r'[^a-zA-Z0-9_-]', '-', company).strip('-').lower() if company else "aicc-poc"
        project_name = project_name or "aicc-poc"

        result = package_and_upload_assets(
            session_id=session_id,
            project_name=project_name,
            include_readme=True,
        )

        if result.get("success"):
            await safe_send_json(websocket, {
                "type": "download_ready",
                "downloadUrl": result.get("download_url", ""),
                "expiresAt": result.get("expires_at", ""),
                "s3Key": result.get("s3_key", ""),
                "fileCount": result.get("file_count", 0),
                "totalSizeMb": result.get("total_size_mb", 0),
            })
        else:
            await safe_send_json(websocket, {
                "type": "error",
                "content": result.get("error", "Asset packaging failed"),
            })
    except Exception as e:
        logger.error(f"Asset packaging error: {e}")
        await safe_send_json(websocket, {"type": "error", "content": f"Asset packaging failed: {e}"})


async def handle_get_progress_ws(websocket: WebSocket, session_id: str):
    """Return session progress from NFS (persistent) + in-memory state."""
    progress_state = _get_frontend_progress(session_id)
    phase = _detect_phase(session_id)

    # In-memory data (may be empty after container restart)
    message_count = 0
    document_mode = False
    if session_id in session_store:
        session = session_store[session_id]
        message_count = len(session.get("conversation_history", []))
        document_mode = session.get("document_mode", False)

    await safe_send_json(websocket, {
        "type": "progress",
        "data": {
            "message_count": message_count,
            "document_mode": document_mode,
            "progressState": progress_state,
            "phase": phase,
        },
    })


async def handle_inject_history_ws(websocket: WebSocket, session_id: str, data: Dict[str, Any]):
    """Restore session from conversation history.

    Converts the DynamoDB-sourced history into Strands format and stores it
    in session["conversation_history"] so the next agent invocation has full
    context.  Previously this function formatted history but never stored it,
    causing the agent to "forget" everything after a reconnect.
    """
    history = data.get("history", [])
    original_session_id = data.get("originalSessionId", session_id)
    session_context = data.get("sessionContext", {})

    session = get_or_create_session(session_id)
    session["session_data"]["original_session_id"] = original_session_id

    # Merge session context fields (companyName, industry, etc.)
    if session_context:
        for key in ("companyName", "industry", "operations", "dbConnected"):
            if key in session_context:
                session["session_data"][key] = session_context[key]

    # Warm the NFS view from S3 before anything reads /mnt/s3 for this session.
    # After an ECS restart, mountpoint-s3 imports directories lazily on first
    # access, so os.scandir on sessions/{sid}/assets/ can return an incomplete
    # tree while imports are in flight. This forces every ancestor dir to
    # materialize and every asset file to stat successfully.
    try:
        from tools.s3_asset_storage import hydrate_session_workspace
        hydrate_session_workspace(original_session_id)
    except Exception as e:
        logger.warning(f"[injectHistory] hydrate failed for {original_session_id}: {e}")

    # Restore workspace
    try:
        from tools.project_workspace import set_workspace_session_id
        ws = set_workspace_session_id(original_session_id)
        restore_result = ws.restore_all()
        workspace_summary = restore_result.get("summary", "")
    except Exception as e:
        logger.warning(f"Workspace restore failed: {e}")
        workspace_summary = ""

    # Restore spec manager cache
    try:
        from tools.spec_manager import restore_specs_from_workspace, restore_flow_config_from_workspace, get_infrastructure_spec
        restore_specs_from_workspace()
        restore_flow_config_from_workspace()
        get_infrastructure_spec()  # triggers lazy restore from NFS/S3
    except Exception:
        pass

    # ── Restore conversation_history into the session ──
    # The injected history from DynamoDB may be richer (more recent) than
    # what NFS had at session creation time.  Convert to Strands format and
    # replace session history if the injected version is newer/larger.
    if history:
        injected_strands = []
        for msg in history:
            role = msg.get("role")
            content = msg.get("content")
            # Bedrock ConverseStream only accepts "user" and "assistant" roles.
            # Drop "system" messages — they cause ValidationException.
            if role not in ("user", "assistant"):
                if role == "system":
                    logger.info(f"[injectHistory] Dropping system-role message (not supported by Bedrock)")
                continue
            if isinstance(content, str):
                injected_strands.append({"role": role, "content": [{"text": content}]})
            elif isinstance(content, list):
                injected_strands.append({"role": role, "content": content})
            # else: skip malformed

        existing_len = len(session.get("conversation_history", []))
        if len(injected_strands) >= existing_len:
            # DynamoDB version is same size or larger — use it
            session["conversation_history"] = _validate_tool_pairs(injected_strands)
            logger.info(
                f"[injectHistory] Replaced conversation_history for {session_id}: "
                f"{existing_len} → {len(session['conversation_history'])} messages"
            )
        else:
            logger.info(
                f"[injectHistory] Kept NFS history for {session_id} "
                f"(NFS={existing_len}, injected={len(injected_strands)})"
            )

        # Persist merged history to NFS so it survives future reconnects
        try:
            _context_store.save_conversation_history(session_id, session["conversation_history"])
        except Exception as e:
            logger.warning(f"[injectHistory] NFS save failed for {session_id}: {e}")

    # Format summary for logging (workspace_summary is still useful for debug)
    logger.info(
        f"[injectHistory] session={session_id} original={original_session_id} "
        f"history_msgs={len(history)} workspace={'yes' if workspace_summary else 'no'}"
    )

    # Build NFS progress for frontend restoration
    _hi_sid = original_session_id or session_id
    _hi_progress = _get_frontend_progress(_hi_sid)

    await safe_send_json(websocket, {
        "type": "history_injected",
        "sessionId": session_id,
        "originalSessionId": original_session_id,
        "messageCount": len(session.get("conversation_history", [])),
        "hasWorkspace": bool(workspace_summary),
        "phase": _detect_phase(_hi_sid),
        "progressState": _hi_progress if _hi_progress else None,
    })


async def handle_create_new_session_ws(websocket: WebSocket, session_id: str):
    """Create a fresh session."""
    if session_id in session_store:
        # Flush before clearing
        try:
            history = session_store[session_id].get("conversation_history", [])
            if history:
                _context_store.save_conversation_history(session_id, history)
        except Exception:
            pass
        del session_store[session_id]

    # Clear sub-agent caches
    clear_research_session(session_id)
    clear_faq_session(session_id)

    # Reset spec manager
    try:
        from tools.spec_manager import clear_all_specs
        clear_all_specs()
    except Exception:
        pass

    # Purge all per-session state from session_context dicts (workspaces,
    # fragment/schema registries, operation specs bucket, etc.).  Uses the
    # effective session id if we had one recorded, otherwise the raw id.
    try:
        from tools.session_context import cleanup_session
        _eff = session_store.get(session_id, {}).get("session_data", {}).get("original_session_id", session_id) \
            if isinstance(session_store.get(session_id), dict) else session_id
        cleanup_session(session_id)
        if _eff and _eff != session_id:
            cleanup_session(_eff)
    except Exception as _ce:
        logger.warning(f"[createNewSession] cleanup_session failed: {_ce}")

    await safe_send_json(websocket, {
        "type": "session_created",
        "sessionId": session_id,
        "phase": _detect_phase(session_id),
    })


# ========================================
# Parallel Sub-Agent Execution (Improvement A)
# ========================================
async def execute_parallel_tools(tool_calls: list, session_id: str) -> list:
    """Execute independent sub-agent calls in parallel using asyncio.gather."""
    async def run_tool(tool_fn, *args, **kwargs):
        return await asyncio.to_thread(tool_fn, *args, **kwargs)

    tasks = [run_tool(fn, *args, **kwargs) for fn, args, kwargs in tool_calls]
    results = await asyncio.gather(*tasks, return_exceptions=True)
    return results
